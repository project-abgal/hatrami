from docx import Document
from docx.shared import RGBColor

def separator_check(characters, idx) -> bool:
    if idx == len(characters) - 1:
        return False

    if characters[idx+1]['number'] is None \
        and characters[idx+1]['reading'] in (']', ']]', '>', '>>'):
        return False

    if characters[idx]['number'] is None \
        and characters[idx]['reading'] in ('[', '[[', '<', '<<'):
        return False

    return True

def output_docx(result, path):
    document = Document()

    def render_characters(sub_word, separator, character_type):
        for i, t in enumerate(sub_word['transliteration']):
            r = p.add_run(t['reading'].replace('[[', '⸢').replace(']]', '⸣'))

            color = RGBColor(0x2d, 0x37, 0x48)
            if t['number'] is None:
                # Annotation
                r.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
            elif character_type == 'hittite':
                r.italic = True
                r.font.color.rgb = color
            elif character_type == 'akkadogram':
                r.italic = True
                color = RGBColor(0xdd, 0x6b, 0x20)
                # color = RGBColor(0xf8, 0xd3, 0x8d)
                r.font.color.rgb = color
            elif character_type == 'sumerogram':
                color = RGBColor(0x38, 0xa1, 0x69)
                # color = RGBColor(0x9a, 0xb6, 0xe4)
                r.font.color.rgb = color
            elif character_type == 'determination':
                r.font.superscript = True
                color = RGBColor(0x00, 0xb5, 0xd8)
                # color = RGBColor(0x9d, 0xec, 0xf9)
                r.font.color.rgb = color

            if t['number'] is not None and t['number'] > 4:
                r = p.add_run(str(t['number']))
                r.font.subscript = True
                r.font.color.rgb = color

            if separator_check(sub_word['transliteration'], i):
                p.add_run(separator)

    def render_subword(sub_word):
        if sub_word['type'] == 'hittite':
            render_characters(sub_word, '-', sub_word['type'])
        elif sub_word['type'] == 'akkadogram':
            render_characters(sub_word, '-', sub_word['type'])
        elif sub_word['type'] == 'sumerogram':
            render_characters(sub_word, '.', sub_word['type'])
        elif sub_word['type'] == 'determination':
            render_characters(sub_word, '.', sub_word['type'])

    for sentence in result:
        if sentence is None:
            continue

        p = document.add_paragraph('', style='List Number')
        for word in sentence:
            for j, sub_word in enumerate(word):
                render_subword(sub_word)

                if j < len(word) - 1:
                    if word[j+1]['type'] == 'determination':
                        continue
                    if j == 0 and word[j]['type'] == 'determination':
                        continue
                    p.add_run('-')

            p.add_run(' ')

    document.save(path)
